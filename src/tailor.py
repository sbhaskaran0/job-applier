"""Per-job resume & cover-letter tailoring (JOB-6).

On-demand generation of a bespoke resume + cover letter for a single posting,
stored under `resumes/<job-slug>/`. Consistent with the project's core
principle, this module holds only the **mechanical** operations — read the base
`.docx`, apply the edits Claude decides on, export to PDF, gather the user's own
cover-letter exemplars, and persist/resolve artifacts. The *reasoning* (which
bullets to re-emphasize, writing the cover letter in the user's voice) is done
by Claude Code via the `/tailor-application` skill, which then calls these tools
to persist the result. No LLM API key lives here.

Storage layout, per job:
    resumes/<job-slug>/
        resume.docx        tailored copy of the base template (formatting kept)
        resume.pdf         exported for upload where the ATS wants a PDF
        cover_letter.txt   the cover letter text (for a free-text field)
        cover_letter.pdf   the cover letter as a file (for a file input)

Job identity reuses `data._application_key` (the (company, role[, url]) tuple
applications.json dedupes on) so tailoring, storage, and the JOB-5 tracker all
share ONE identifier.
"""

import hashlib
import re

from . import config, data


# --------------------------------------------------------------------------- #
# Job identity → filesystem-safe slug (shares data._application_key)
# --------------------------------------------------------------------------- #
def job_slug(company: str = "", job_title: str = "", url: str = "") -> str:
    """A filesystem-safe folder name for this job, derived from the SAME identity
    applications.json dedupes on (`data._application_key`). Normalized
    `company-job_title`; if the title/company are missing, a short hash of the
    URL keeps it unique."""
    company_n, title_n, url_part = data._application_key(company, job_title, url)
    base = "-".join(p for p in (company_n, title_n) if p).strip()
    slug = re.sub(r"[^a-z0-9]+", "-", base).strip("-")
    if not slug:
        # No usable company/title — fall back to a stable hash of the url tuple.
        h = hashlib.sha1((url_part or url or "").encode("utf-8")).hexdigest()[:12]
        slug = f"job-{h}" if h else "job-unknown"
    return slug[:120]


def job_dir(company: str = "", job_title: str = "", url: str = ""):
    """Path to resumes/<job-slug>/ for this job (not created)."""
    return config.RESUMES_DIR / job_slug(company, job_title, url)


# --------------------------------------------------------------------------- #
# Base .docx template — read structure so Claude can plan edits
# --------------------------------------------------------------------------- #
def _iter_paragraphs(parent):
    """Yield every paragraph in document order, descending into tables (resume
    templates often lay content out in tables). Order is stable, so the position
    in this stream is a paragraph's addressable index."""
    from docx.document import Document as _Doc
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Doc):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)


def _open_base_document():
    """Open the user's base resume.docx. Raises a clear error if absent."""
    from docx import Document

    base = config.base_resume_docx()
    if base is None:
        raise FileNotFoundError(
            "No resume.docx base template found. Per-job resume tailoring edits a "
            "`resume.docx` you provide (drop it in the project root, alongside "
            "resume.pdf). Until then, applications use the default resume."
        )
    return Document(str(base))


def read_resume_template() -> dict:
    """Return the base resume as an ordered, indexed list of paragraphs so Claude
    can decide which to rewrite/reorder/drop. Each item:
      {index, text, style, is_bullet, is_empty}
    `index` is the stable address to pass back in `tailor_resume` edits."""
    doc = _open_base_document()
    paras = []
    for i, p in enumerate(_iter_paragraphs(doc)):
        style = (p.style.name if p.style is not None else "") or ""
        text = p.text or ""
        is_bullet = ("list" in style.lower()
                     or bool(p._p.find(
                         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr/"
                         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")))
        paras.append({"index": i, "text": text, "style": style,
                      "is_bullet": is_bullet, "is_empty": not text.strip()})
    return {"base": str(config.base_resume_docx()), "paragraph_count": len(paras),
            "paragraphs": paras}


# --------------------------------------------------------------------------- #
# Apply edits → tailored .docx (+ PDF)
# --------------------------------------------------------------------------- #
def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace a paragraph's text while KEEPING its formatting: write into the
    first run (inheriting its font/bold/size) and blank the remaining runs. If
    the paragraph has no runs, add one (inherits the paragraph/style default)."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _delete_paragraph(paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)
    el._p = el._element = None


def tailor_resume(company: str = "", job_title: str = "", url: str = "",
                  edits: list | None = None) -> dict:
    """Apply `edits` to a COPY of the base resume.docx and save it under
    resumes/<job-slug>/resume.docx, then export a resume.pdf for upload.

    `edits` is a list of ops addressing paragraph `index` from
    read_resume_template (the base is re-read here, so indices line up):
      {"op": "replace", "index": N, "text": "..."}  — rewrite a paragraph's text,
                                                       preserving its formatting.
      {"op": "delete",  "index": N}                 — remove a paragraph.
    Reorder/re-emphasize bullets by `replace`-ing their text in the desired
    order and `delete`-ing the least JD-relevant ones. Formatting/layout of the
    template is preserved (no text-to-PDF regeneration).

    Returns {slug, dir, docx_path, pdf_path, pdf_exported, edits_applied, note}.
    """
    edits = edits or []
    doc = _open_base_document()
    paras = list(_iter_paragraphs(doc))
    n = len(paras)

    applied, errors = 0, []
    # Apply against original indices; deletes don't reindex the others because we
    # hold the Paragraph objects directly.
    for e in edits:
        op = (e.get("op") or "").lower()
        idx = e.get("index")
        if not isinstance(idx, int) or not (0 <= idx < n):
            errors.append(f"index {idx} out of range 0..{n - 1}")
            continue
        p = paras[idx]
        if p._element is None:
            errors.append(f"index {idx} already deleted")
            continue
        if op == "replace":
            _set_paragraph_text(p, e.get("text", ""))
            applied += 1
        elif op == "delete":
            _delete_paragraph(p)
            applied += 1
        else:
            errors.append(f"index {idx}: unknown op {op!r}")

    out_dir = job_dir(company, job_title, url)
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / "resume.docx"
    doc.save(str(docx_path))

    pdf_path = out_dir / "resume.pdf"
    pdf_ok, pdf_note = _export_pdf(docx_path, pdf_path)

    return {"slug": out_dir.name, "dir": str(out_dir),
            "docx_path": str(docx_path),
            "pdf_path": str(pdf_path) if pdf_ok else None,
            "pdf_exported": pdf_ok, "edits_applied": applied,
            "edit_errors": errors,
            "note": pdf_note if not pdf_ok else "tailored resume saved (docx + pdf)"}


# --------------------------------------------------------------------------- #
# Cover letter — exemplars in, artifact out
# --------------------------------------------------------------------------- #
_COVER_PAT = re.compile(r"cover", re.I)
_SAMPLE_PAT = re.compile(r"writing[_ ]sample|personal statement|statement of purpose|essay", re.I)
# Reference material and structured knowledge files that are NOT the user's own
# free-form prose — kept out of the voice corpus. background/stories/preferences
# still feed the cover letter as SUBSTANCE via search_context, separately.
_NON_VOICE_PAT = re.compile(r"transcript|resume|^background\.|^stories\.|^preferences\.", re.I)


def cover_letter_examples() -> dict:
    """Return the FULL text of the user's own writing from context/, as few-shot
    exemplars for matching their VOICE — sentence rhythm, structure, tone — when
    drafting a cover letter. Three buckets, all the user's authentic prose:
      - cover_letters: past cover letters (primary style source).
      - writing_samples: essays / personal statements / writing samples.
      - responses: past answers to application questions (e.g. "Describe a
        product you managed…", "the hardest you've worked…") — the user's own
        first-person responses, useful as BOTH voice and substance.
    The catch-all `responses` bucket also picks up any other free-form .txt/.md
    the user drops into context/, so newly pasted writing is used automatically.
    Structured knowledge (background/stories/preferences) and reference files
    (transcript, resume) are excluded — those feed substance via search_context.
    Each item: {name, type, text}. Empty/no-text files are dropped."""
    covers, samples, responses = [], [], []
    for path in sorted(config.CONTEXT_DIR.iterdir()):
        if not path.is_file():
            continue
        name = path.name
        if _NON_VOICE_PAT.search(name):
            continue
        is_cover = bool(_COVER_PAT.search(name))
        is_sample = bool(_SAMPLE_PAT.search(name))
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            # Only pull named cover letters / writing samples from PDFs — avoid
            # slurping arbitrary reference PDFs as if they were the user's voice.
            if not (is_cover or is_sample):
                continue
            text = config.extract_pdf_text(path)
        elif suffix in (".txt", ".md"):
            text = path.read_text(encoding="utf-8", errors="ignore")
        else:
            continue
        if not text.strip():
            continue
        item = {"name": name, "text": text.strip()}
        if is_cover:
            item["type"] = "cover_letter"
            covers.append(item)
        elif is_sample:
            item["type"] = "writing_sample"
            samples.append(item)
        else:
            item["type"] = "response"   # application-question answers + other prose
            responses.append(item)
    return {"cover_letters": covers, "writing_samples": samples,
            "responses": responses,
            "count": len(covers) + len(samples) + len(responses)}


def _text_to_docx(text: str, path) -> None:
    """Write plain text (blank-line-separated paragraphs) to a simple .docx."""
    from docx import Document

    doc = Document()
    for block in re.split(r"\n\s*\n", text.strip()):
        doc.add_paragraph(block.strip())
    doc.save(str(path))


def save_cover_letter(company: str = "", job_title: str = "", url: str = "",
                      text: str = "") -> dict:
    """Persist a drafted cover letter under resumes/<job-slug>/: cover_letter.txt
    (for a free-text field) and cover_letter.pdf (for a file input). Returns
    {slug, dir, txt_path, pdf_path, pdf_exported, note}."""
    if not text.strip():
        raise ValueError("cover letter text is empty")
    out_dir = job_dir(company, job_title, url)
    out_dir.mkdir(parents=True, exist_ok=True)
    txt_path = out_dir / "cover_letter.txt"
    txt_path.write_text(text.strip() + "\n", encoding="utf-8")

    docx_path = out_dir / "cover_letter.docx"
    pdf_path = out_dir / "cover_letter.pdf"
    try:
        _text_to_docx(text, docx_path)
        pdf_ok, pdf_note = _export_pdf(docx_path, pdf_path)
    except Exception as exc:  # noqa: BLE001 — never lose the .txt over a PDF issue
        pdf_ok, pdf_note = False, f"pdf export failed: {type(exc).__name__}: {exc}"

    return {"slug": out_dir.name, "dir": str(out_dir), "txt_path": str(txt_path),
            "pdf_path": str(pdf_path) if pdf_ok else None,
            "pdf_exported": pdf_ok,
            "note": "cover letter saved (txt + pdf)" if pdf_ok
                    else f"cover letter saved (txt only) — {pdf_note}"}


# --------------------------------------------------------------------------- #
# Apply-time resolution: tailored artifact if present, else default
# --------------------------------------------------------------------------- #
def job_artifacts(company: str = "", job_title: str = "", url: str = "") -> dict:
    """Resolve which resume/cover-letter to use for a job at apply time.

    Checks resumes/<job-slug>/ for tailored artifacts and FALLS BACK to the
    default resume when none exist. Returns:
      {slug, dir, tailored (bool), resume_path, resume_is_tailored,
       cover_letter_path, cover_letter_text, has_cover_letter}
    `resume_path` is always usable (tailored PDF/DOCX if present, else the
    project default from config.resume_upload_path)."""
    out_dir = job_dir(company, job_title, url)
    resume_tailored = None
    for name in ("resume.pdf", "resume.docx"):
        p = out_dir / name
        if p.exists():
            resume_tailored = p
            break
    cover_file = None
    for name in ("cover_letter.pdf", "cover_letter.docx"):
        p = out_dir / name
        if p.exists():
            cover_file = p
            break
    cover_txt = out_dir / "cover_letter.txt"
    cover_text = cover_txt.read_text(encoding="utf-8") if cover_txt.exists() else ""

    resume_path = resume_tailored or config.resume_upload_path()
    return {"slug": out_dir.name, "dir": str(out_dir),
            "tailored": bool(resume_tailored or cover_file or cover_text),
            "resume_path": str(resume_path),
            "resume_is_tailored": resume_tailored is not None,
            "cover_letter_path": str(cover_file) if cover_file else None,
            "cover_letter_text": cover_text,
            "has_cover_letter": bool(cover_file or cover_text)}


# --------------------------------------------------------------------------- #
# DOCX -> PDF export (Word COM on Windows; graceful if unavailable)
# --------------------------------------------------------------------------- #
_WD_FORMAT_PDF = 17  # Word's wdFormatPDF SaveAs code


def _convert_via_word(docx_path, pdf_path) -> tuple[bool, str]:
    """Drive Microsoft Word directly (win32com) in a DEDICATED Word process
    (DispatchEx) — one open→SaveAs→close→quit per call. A fresh process each
    time sidesteps docx2pdf's intermittent `Open.SaveAs` failure when a Word
    instance is reused across conversions."""
    try:
        import pythoncom  # type: ignore
        import win32com.client as win32  # type: ignore
    except Exception as exc:  # noqa: BLE001 — pywin32 absent
        return False, f"win32com unavailable ({type(exc).__name__})"
    word = None
    coinit = False
    try:
        pythoncom.CoInitialize()
        coinit = True
        word = win32.DispatchEx("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=1)
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=_WD_FORMAT_PDF)
            doc.Close(False)
        finally:
            word.Quit()
    except Exception as exc:  # noqa: BLE001
        return False, f"{type(exc).__name__}: {exc}"
    finally:
        if coinit:
            pythoncom.CoUninitialize()
    return (pdf_path.exists(), "" if pdf_path.exists() else "no file written")


def _convert_via_docx2pdf(docx_path, pdf_path) -> tuple[bool, str]:
    """docx2pdf drives Microsoft Word — Windows via COM, macOS via AppleScript.
    So this covers a Mac that HAS Word (the win32com path above is Windows-only)."""
    try:
        from docx2pdf import convert
    except Exception as exc:  # noqa: BLE001
        return False, f"docx2pdf unavailable ({type(exc).__name__})"
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as exc:  # noqa: BLE001
        return False, f"{type(exc).__name__}: {exc}"
    return (pdf_path.exists(), "" if pdf_path.exists() else "no file written")


def _find_soffice():
    """Locate a LibreOffice/OpenOffice `soffice` binary (cross-platform, no MS
    Word needed). Checks PATH then the usual per-OS install locations."""
    import shutil

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
        "/usr/bin/soffice", "/usr/local/bin/soffice",            # Linux
        "/opt/libreoffice/program/soffice",
    ]
    for c in candidates:
        from pathlib import Path
        if Path(c).exists():
            return c
    return None


def _convert_via_libreoffice(docx_path, pdf_path) -> tuple[bool, str]:
    """Convert with LibreOffice headless — works on Windows/macOS/Linux with NO
    Microsoft Word. `soffice` writes <stem>.pdf into --outdir; we rename if the
    caller wanted a different name."""
    import subprocess

    soffice = _find_soffice()
    if not soffice:
        return False, "LibreOffice (soffice) not found"
    outdir = pdf_path.parent
    try:
        proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir",
             str(outdir), str(docx_path.resolve())],
            capture_output=True, text=True, timeout=120)
    except Exception as exc:  # noqa: BLE001
        return False, f"soffice failed ({type(exc).__name__}: {exc})"
    produced = outdir / (docx_path.stem + ".pdf")
    if produced != pdf_path and produced.exists():
        produced.replace(pdf_path)
    if pdf_path.exists():
        return True, ""
    return False, (proc.stderr or proc.stdout or "no file written").strip()[:200]


def _export_pdf(docx_path, pdf_path) -> tuple[bool, str]:
    """Export a .docx to .pdf, trying converters in fidelity order and stopping
    at the first that works — cross-platform:
      1. MS Word via win32com COM (Windows + Word; retried once), then
      2. docx2pdf (Windows COM OR macOS AppleScript — a Mac WITH Word), then
      3. LibreOffice headless (`soffice`; Windows/macOS/Linux, NO Word needed).
    Returns (ok, note) and never raises — if none is available the caller still
    has the .docx/.txt and the note explains how to convert manually."""
    if pdf_path.exists():
        try:
            pdf_path.unlink()  # stale output would read as a false success
        except OSError:
            pass
    notes = []
    for attempt in (1, 2):  # Word COM: high fidelity, retry the known flaky call
        ok, note = _convert_via_word(docx_path, pdf_path)
        if ok:
            return True, "pdf exported (Word)" + (" [retry]" if attempt == 2 else "")
        notes.append(note)
    for fn, label in ((_convert_via_docx2pdf, "docx2pdf"),
                      (_convert_via_libreoffice, "LibreOffice")):
        ok, note = fn(docx_path, pdf_path)
        if ok:
            return True, f"pdf exported ({label})"
        notes.append(note)
    detail = "; ".join(n for n in notes if n) or "no converter available"
    return False, (f"PDF export failed ({detail}). docx saved — install MS Word "
                   "or LibreOffice, or open the .docx and Save As PDF.")
